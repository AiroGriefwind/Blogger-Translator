from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class DocxFormatter:
    def build(
        self,
        output_path: str | Path,
        title_en: str,
        header_byline_en: str,
        body_blocks: list[str],
        ending_author_en: str,
        ending_column_en: str,
        captions_blocks: list[str],
    ) -> Path:
        doc = Document()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_en)
        title_run.bold = True
        title_run.font.size = Pt(15)

        if header_byline_en.strip():
            author = doc.add_paragraph()
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author.add_run(header_byline_en)
            author_run.italic = True
            author_run.font.size = Pt(15)

        for block in body_blocks:
            if str(block).strip():
                p = doc.add_paragraph()
                run = p.add_run(block)
                run.font.size = Pt(13)
                p.paragraph_format.line_spacing = 2.0
            else:
                doc.add_paragraph("")

        if ending_author_en.strip():
            end_author = doc.add_paragraph()
            end_author.alignment = WD_ALIGN_PARAGRAPH.LEFT
            end_author_run = end_author.add_run(ending_author_en)
            end_author_run.italic = True
            end_author_run.font.size = Pt(13)
            end_author.paragraph_format.line_spacing = 2.0

        if ending_column_en.strip():
            end_column = doc.add_paragraph()
            end_column.alignment = WD_ALIGN_PARAGRAPH.LEFT
            end_column_run = end_column.add_run(ending_column_en)
            end_column_run.font.size = Pt(13)
            end_column.paragraph_format.line_spacing = 2.0

        doc.add_paragraph("")
        cap_title = doc.add_paragraph()
        cap_title_run = cap_title.add_run("Captions:")
        cap_title_run.bold = True
        cap_title.paragraph_format.line_spacing = 2.0

        for cap in captions_blocks:
            p = doc.add_paragraph()
            run = p.add_run(cap)
            run.font.size = Pt(13)
            p.paragraph_format.line_spacing = 2.0

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        return output

